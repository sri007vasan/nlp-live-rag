"""
Document parser module for handling multiple file formats.
Supports PDF, DOCX (Word), and XLSX (Excel) files.
"""

import io
import logging
import os
from pathlib import Path
from typing import Optional, List, Tuple

# PDF parsing
try:
    import pypdf
    PDF_AVAILABLE = True
except ImportError:
    PDF_AVAILABLE = False
    logging.warning("pypdf not installed. PDF parsing will not be available.")

# DOCX parsing
try:
    from docx import Document as DocxDocument
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    logging.warning("python-docx not installed. DOCX parsing will not be available.")

# XLSX parsing
try:
    import openpyxl
    XLSX_AVAILABLE = True
except ImportError:
    XLSX_AVAILABLE = False
    logging.warning("openpyxl not installed. XLSX parsing will not be available.")


class DocumentParser:
    """Universal document parser supporting PDF, DOCX, and XLSX formats."""

    @staticmethod
    def parse_pdf(file_path: str) -> Tuple[str, bool]:
        """
        Parse PDF file and extract text.
        
        Args:
            file_path: Path to the PDF file
            
        Returns:
            Tuple of (extracted_text, success)
        """
        if not PDF_AVAILABLE:
            return "", False
            
        try:
            text = ""
            with open(file_path, 'rb') as file:
                pdf_reader = pypdf.PdfReader(file)
                for page in pdf_reader.pages:
                    text += page.extract_text()
            return text, True
        except Exception as e:
            logging.error(f"Error parsing PDF {file_path}: {str(e)}")
            return "", False

    @staticmethod
    def parse_docx(file_path: str) -> Tuple[str, bool]:
        """
        Parse DOCX (Word) file and extract text.
        
        Args:
            file_path: Path to the DOCX file
            
        Returns:
            Tuple of (extracted_text, success)
        """
        if not DOCX_AVAILABLE:
            return "", False
            
        try:
            doc = DocxDocument(file_path)
            text = ""
            
            # Extract paragraphs
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
            
            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = " | ".join(cell.text for cell in row.cells)
                    text += row_text + "\n"
            
            return text, True
        except Exception as e:
            logging.error(f"Error parsing DOCX {file_path}: {str(e)}")
            return "", False

    @staticmethod
    def parse_xlsx(file_path: str) -> Tuple[str, bool]:
        """
        Parse XLSX (Excel) file and extract text.
        
        Args:
            file_path: Path to the XLSX file
            
        Returns:
            Tuple of (extracted_text, success)
        """
        if not XLSX_AVAILABLE:
            return "", False
            
        try:
            workbook = openpyxl.load_workbook(file_path)
            text = ""
            
            for sheet_name in workbook.sheetnames:
                sheet = workbook[sheet_name]
                text += f"\n=== Sheet: {sheet_name} ===\n"
                
                for row in sheet.iter_rows(values_only=True):
                    row_text = " | ".join(str(cell) if cell is not None else "" for cell in row)
                    text += row_text + "\n"
            
            return text, True
        except Exception as e:
            logging.error(f"Error parsing XLSX {file_path}: {str(e)}")
            return "", False

    @staticmethod
    def parse_document(file_path: str) -> Tuple[str, bool]:
        """
        Parse any supported document format based on file extension.
        
        Args:
            file_path: Path to the document file
            
        Returns:
            Tuple of (extracted_text, success)
        """
        file_path = str(file_path)
        file_extension = Path(file_path).suffix.lower()
        
        if file_extension == ".pdf":
            return DocumentParser.parse_pdf(file_path)
        elif file_extension == ".docx":
            return DocumentParser.parse_docx(file_path)
        elif file_extension in [".doc"]:
            # For .doc files, we would need python-docx2docx or similar
            # For now, we'll try to treat as DOCX if possible
            return DocumentParser.parse_docx(file_path)
        elif file_extension == ".xlsx":
            return DocumentParser.parse_xlsx(file_path)
        else:
            logging.warning(f"Unsupported file format: {file_extension}")
            return "", False

    @staticmethod
    def get_supported_formats() -> List[str]:
        """
        Get list of supported file formats based on installed packages.
        
        Returns:
            List of supported file extensions
        """
        formats = []
        if PDF_AVAILABLE:
            formats.append(".pdf")
        if DOCX_AVAILABLE:
            formats.extend([".docx", ".doc"])
        if XLSX_AVAILABLE:
            formats.append(".xlsx")
        return formats

    @staticmethod
    def is_format_supported(file_extension: str) -> bool:
        """
        Check if a file format is supported.
        
        Args:
            file_extension: File extension (e.g., ".pdf", ".docx")
            
        Returns:
            True if format is supported, False otherwise
        """
        return file_extension.lower() in DocumentParser.get_supported_formats()


class DocumentParserFactory:
    """Factory class for creating parsers for specific document types."""
    
    _parsers = {
        ".pdf": DocumentParser.parse_pdf,
        ".docx": DocumentParser.parse_docx,
        ".doc": DocumentParser.parse_docx,  # Treat as DOCX
        ".xlsx": DocumentParser.parse_xlsx,
    }
    
    @classmethod
    def get_parser(cls, file_extension: str):
        """
        Get parser for a specific file extension.
        
        Args:
            file_extension: File extension (e.g., ".pdf")
            
        Returns:
            Parser function or None if not supported
        """
        return cls._parsers.get(file_extension.lower())
    
    @classmethod
    def register_parser(cls, file_extension: str, parser_func):
        """
        Register a custom parser for a file extension.
        
        Args:
            file_extension: File extension (e.g., ".custom")
            parser_func: Parser function that takes file_path and returns (text, success)
        """
        cls._parsers[file_extension.lower()] = parser_func
